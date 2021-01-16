# GUIWiki.py
import wikipedia

#python to docx
from docx import Document 
import wikipedia

#กดtab มันจะเข้ามาอยู่ใน function def Wiki

def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)

    #summary บทความ
    data = wikipedia.summary(keyword)

    #page + content บทความทั้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content

    doc = Document() #สร้างไฟล์เวิร์ดใน python
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2)
    doc.save(keyword+'.docx')
    print('สร้างไฟล์สำเร็จ')

#เปลี่ยนภาษา
wikipedia.set_lang('th')
from tkinter import *
from tkinter import ttk
from tkinter import messagebox

GUI = Tk()
GUI.title('โปรแกรม wiki by FJen')
GUI.geometry('400x300')

#config
FONT1 = ('Angsana New', 15)

#คำอธิบาย
L= ttk.Label(GUI, text='ค้นหาบทความ', font= FONT1)
L.pack()

# ช่องค้นหาข้อมูล
v_search = StringVar()
E1 = ttk.Entry(GUI,textvariable = v_search,font=FONT1,width = 40)
E1.pack(pady = 10)

#Buttonค้นหา
def Search():
    keyword = v_search.get() #.get คือ ดึงข้อมูลเข้ามาได้แค้ StringVar เท่านั้น
    try:
        #ลองค้นหาดูว่าได้ผลลัพธ์หรือไม่ หากได้ให้ผ่านไป
        language = v_radio.get()
        Wiki(keyword,language)
        messagebox.showinfo('บันทึกสำเร็จ','ค้นหาข้อความสำเร็จ บันทึกเรียบร้อยแล้ว')
    except:
        #หากรันคำสั่ง Error แล้วมีปัญหา แสดงข้อความแจ้งเตือน
        messagebox.showwarning('Keyword Error!', 'กรุณากรอกคำค้นหาใหม่')
        
        
    #result = wikipedia.summary(keyword)
    #print(result)

B1 = ttk.Button(GUI,text = 'Search', command = Search)
B1.pack(ipadx=20,ipady=10)

# เลือกภาษา
F1 = Frame(GUI)
F1.pack(pady = 20)

v_radio = StringVar()# ช่องเก็บข้อมูลภาษา

RB1 = ttk.Radiobutton(F1,text='ภาษาไทย',variable = v_radio, value='th')
RB2 = ttk.Radiobutton(F1,text='อังกฤษ',variable = v_radio, value='en')
RB3 = ttk.Radiobutton(F1,text='จีน',variable = v_radio, value='zh')
RB1.invoke() #สั่งให้ค่าเริ่มต้นเป็นภาษาไทย


RB1.grid(row=0,column=0) #pack จะแสดงผลเป็นบรรทัดๆลงมา แต่ grid จะแสดงผลไปทางขวาเราจึงต้องใส่ row กับ column
RB2.grid(row=0,column=1)
RB3.grid(row=0,column=2)

GUI.mainloop()
