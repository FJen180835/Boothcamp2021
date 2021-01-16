#  test-docx.py
'''
from กับ import ต่างกันยังไง
from มี class ข้างในคือ document และจะสามารถใช้ได้คำสั่งเดียวคือ document อีกเลยต้อง from
ส่วน import คือ ดึงมาหมดเลย ใช้งานได้ทุกคำสั่ง เช่น set_lang, summary
'''

from docx import Document 
import wikipedia



#กดtab มันจะเข้ามาอยู่ใน function def Wiki

def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)

    #summary บทความ
    data = wikipedia.summary(keyword)

    #page + content บทความทั้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content

    doc = Document() #สร้างไฟล์เวิร์ดใน python
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2)
    doc.save(keyword+'.docx')
    print('สร้างไฟล์สำเร็จ')

Wiki('ประเทศไทย')
Wiki('United state','en')
Wiki('ประเทศญี่ปุ่น','jp')
#ก่อนรันไฟล์ให้ปิดโปรแกรม word ก่อ
